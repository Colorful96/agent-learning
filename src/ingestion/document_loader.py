from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}


def extract_text(file_path: str) -> str:
    """根据文件扩展名提取文本。"""

    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("只支持 TXT、DOCX 和 PDF 文件。")

    if suffix == ".txt":
        text = path.read_text(encoding="utf-8")

    elif suffix == ".docx":
        document = Document(str(path))
        parts = []

        # 提取 Word 段落
        for paragraph in document.paragraphs:
            content = paragraph.text.strip()

            if content:
                parts.append(content)

        # 提取 Word 表格
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                if any(cells):
                    parts.append(" | ".join(cells))

        text = "\n\n".join(parts)

    else:
        reader = PdfReader(str(path))
        pages = []

        # 提取 PDF 每一页的文本
        for page in reader.pages:
            page_text = page.extract_text() or ""

            if page_text.strip():
                pages.append(page_text)

        text = "\n\n".join(pages)

    if not text.strip():
        raise ValueError("没有从文件中提取到有效文本。" "扫描版 PDF 可能需要 OCR。")

    return text
